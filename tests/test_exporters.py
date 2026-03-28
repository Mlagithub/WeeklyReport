"""
Test scaffolding for exporters module.

These tests define the expected behavior for ExporterBase, ExporterFactory,
ImageResolver, and dependency imports. Tests are designed to FAIL initially
and will pass once implementation is complete in Wave 1.

Phase: 08-export-foundation
Plan: 00 (Test Scaffolding)
"""

from io import BytesIO
from typing import Any, List

import pytest


class TestExporterBase:
    """Test ExporterBase abstract class behavior."""

    def test_export_method_exists(self):
        """Verify ExporterBase has export(records, **options) method."""
        from exporters.base import ExporterBase
        assert hasattr(ExporterBase, 'export')
        assert callable(ExporterBase.export)

    def test_file_extension_property(self):
        """Verify file_extension property exists and is abstract."""
        from exporters.base import ExporterBase
        assert hasattr(ExporterBase, 'file_extension')
        # Verify it's abstract by checking __isabstractmethod__
        assert ExporterBase.file_extension.fget.__isabstractmethod__

    def test_mime_type_property(self):
        """Verify mime_type property exists and is abstract."""
        from exporters.base import ExporterBase
        assert hasattr(ExporterBase, 'mime_type')
        # Verify it's abstract by checking __isabstractmethod__
        assert ExporterBase.mime_type.fget.__isabstractmethod__

    def test_cannot_instantiate_base(self):
        """Verify ExporterBase cannot be instantiated directly."""
        from exporters.base import ExporterBase
        with pytest.raises(TypeError):
            ExporterBase()


class TestExporterFactory:
    """Test ExporterFactory class behavior."""

    def test_register_exporter(self):
        """Verify register(format, exporter_class) class method exists."""
        from exporters import ExporterBase, ExporterFactory

        # Create a concrete exporter for testing
        class MockExporter(ExporterBase):
            @property
            def file_extension(self):
                return 'mock'

            @property
            def mime_type(self):
                return 'application/mock'

            def _generate(self, data, options):
                from io import BytesIO
                return BytesIO(b'mock')

        # Test that register method exists and works
        ExporterFactory.register('mock', MockExporter)
        assert 'mock' in ExporterFactory.supported_formats()

        # Clean up registry
        ExporterFactory._registry.pop('mock', None)

    def test_get_exporter_returns_instance(self):
        """Verify get_exporter(format) returns exporter instance."""
        from exporters import ExporterBase, ExporterFactory

        # Create a concrete exporter for testing
        class TestExporter(ExporterBase):
            @property
            def file_extension(self):
                return 'test'

            @property
            def mime_type(self):
                return 'application/test'

            def _generate(self, data, options):
                from io import BytesIO
                return BytesIO(b'test')

        ExporterFactory.register('test', TestExporter)
        exporter = ExporterFactory.get_exporter('test')

        assert isinstance(exporter, ExporterBase)
        assert isinstance(exporter, TestExporter)

        # Clean up registry
        ExporterFactory._registry.pop('test', None)

    def test_get_exporter_invalid_format_raises(self):
        """Verify ValueError raised for unsupported format."""
        from exporters import ExporterFactory

        # Clear registry to ensure no formats registered
        original_registry = ExporterFactory._registry.copy()
        ExporterFactory._registry.clear()

        try:
            with pytest.raises(ValueError) as exc_info:
                ExporterFactory.get_exporter('invalid')

            assert 'Unsupported export format' in str(exc_info.value)
            assert 'invalid' in str(exc_info.value)
        finally:
            # Restore registry
            ExporterFactory._registry.update(original_registry)

    def test_supported_formats_returns_list(self):
        """Verify supported_formats() returns list of format strings."""
        from exporters import ExporterBase, ExporterFactory

        # Create a concrete exporter for testing
        class FormatExporter(ExporterBase):
            @property
            def file_extension(self):
                return 'format'

            @property
            def mime_type(self):
                return 'application/format'

            def _generate(self, data, options):
                from io import BytesIO
                return BytesIO(b'format')

        ExporterFactory.register('format', FormatExporter)
        formats = ExporterFactory.supported_formats()

        assert isinstance(formats, list)
        assert 'format' in formats

        # Clean up registry
        ExporterFactory._registry.pop('format', None)


class TestImageResolver:
    """Test ImageResolver class behavior."""

    def test_resolve_url_files_prefix(self):
        """Verify /files/filename.jpg converts to absolute path."""
        from exporters.image_resolver import ImageResolver
        resolver = ImageResolver('/var/www/uploads')
        result = resolver.resolve_url('/files/test.jpg')
        assert result == '/var/www/uploads/test.jpg'

    def test_resolve_url_external_returns_none(self):
        """Verify external URLs return None."""
        from exporters.image_resolver import ImageResolver
        resolver = ImageResolver('/var/www/uploads')
        assert resolver.resolve_url('http://example.com/img.jpg') is None
        assert resolver.resolve_url('https://example.com/img.jpg') is None

    def test_get_image_bytes_returns_bytes(self):
        """Verify get_image_bytes(url) returns bytes or None."""
        import os
        import tempfile

        from exporters.image_resolver import ImageResolver

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test image file
            test_file = os.path.join(tmpdir, 'test.png')
            test_content = b'\x89PNG\r\n\x1a\n'  # PNG header bytes
            with open(test_file, 'wb') as f:
                f.write(test_content)

            resolver = ImageResolver(tmpdir)
            result = resolver.get_image_bytes('/files/test.png')
            assert result == test_content

            # Non-existent file returns None
            assert resolver.get_image_bytes('/files/missing.png') is None

    def test_image_exists_returns_bool(self):
        """Verify image_exists(url) returns boolean."""
        import os
        import tempfile

        from exporters.image_resolver import ImageResolver

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test file
            test_file = os.path.join(tmpdir, 'exists.jpg')
            with open(test_file, 'wb') as f:
                f.write(b'test')

            resolver = ImageResolver(tmpdir)
            assert resolver.image_exists('/files/exists.jpg') is True
            assert resolver.image_exists('/files/missing.jpg') is False


class TestDependencies:
    """Test that required export dependencies are installed."""

    def test_python_docx_installed(self):
        """Verify python-docx is importable as docx."""
        import docx
        assert docx is not None

    def test_weasyprint_installed(self):
        """Verify weasyprint is importable."""
        import weasyprint
        assert weasyprint is not None

    def test_htmldocx_installed(self):
        """Verify htmldocx is importable."""
        from htmldocx import HtmlToDocx
        assert HtmlToDocx is not None


class TestPdfExporter:
    """Test PdfExporter class behavior.

    Tests for PDF export functionality implemented in Phase 09 Plan 01.
    """

    def test_file_extension(self):
        """Verify PdfExporter.file_extension returns 'pdf'."""
        from exporters.pdf import PdfExporter
        exporter = PdfExporter(uploads_path='/tmp')
        assert exporter.file_extension == 'pdf'

    def test_mime_type(self):
        """Verify PdfExporter.mime_type returns 'application/pdf'."""
        from exporters.pdf import PdfExporter
        exporter = PdfExporter(uploads_path='/tmp')
        assert exporter.mime_type == 'application/pdf'

    def test_export_returns_bytesio(self):
        """Verify export() returns BytesIO with PDF content."""
        from unittest.mock import MagicMock

        from exporters.pdf import PdfExporter

        # Create mock records
        record = MagicMock()
        record.content = '<p>Test content</p>'
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = PdfExporter(uploads_path='/tmp')
        result = exporter.export([record], title='Test Report')

        assert isinstance(result, BytesIO)
        # Verify PDF header
        result.seek(0)
        header = result.read(5)
        assert header == b'%PDF-'

    def test_image_embedding(self):
        """Verify url_fetcher resolves /files/ URLs for image embedding."""
        import os
        import tempfile
        from unittest.mock import MagicMock

        from exporters.pdf import PdfExporter

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test image file
            test_image = os.path.join(tmpdir, 'test.png')
            with open(test_image, 'wb') as f:
                f.write(b'\x89PNG\r\n\x1a\n')  # PNG header

            # Create mock record with image
            record = MagicMock()
            record.content = '<img src="/files/test.png">'
            record.date.strftime = lambda fmt: '2026-03-26'

            exporter = PdfExporter(uploads_path=tmpdir)

            # Test url_fetcher directly
            result = exporter._resolve_image_url('/files/test.png')
            assert result['mime_type'] == 'image/png'
            assert result['string'] == b'\x89PNG\r\n\x1a\n'

    def test_headers_footers(self):
        """Verify CSS Paged Media generates headers with title and footers with page numbers."""
        from unittest.mock import MagicMock

        from exporters.pdf import PdfExporter

        record = MagicMock()
        record.content = '<p>Content</p>'
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = PdfExporter(uploads_path='/tmp')
        html = exporter._build_html([record], title='Test Title', include_date=True)

        # Verify CSS Paged Media elements
        assert '@page' in html
        assert '@top-center' in html
        assert '@bottom-left' in html
        assert '@bottom-right' in html
        assert 'counter(page)' in html
        assert 'running(header)' in html
        assert 'Test Title' in html
        assert 'Generated:' in html


class TestDocxExporter:
    """Test DocxExporter class behavior.

    Tests for DOCX export functionality to be implemented in Phase 10 Plan 01.
    These tests are designed to FAIL initially and will pass once DocxExporter
    is implemented.
    """

    def test_file_extension(self):
        """Verify DocxExporter.file_extension returns 'docx'."""
        from exporters.docx import DocxExporter
        exporter = DocxExporter(uploads_path='/tmp')
        assert exporter.file_extension == 'docx'

    def test_mime_type(self):
        """Verify DocxExporter.mime_type returns correct DOCX MIME type."""
        from exporters.docx import DocxExporter
        exporter = DocxExporter(uploads_path='/tmp')
        expected_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        assert exporter.mime_type == expected_mime

    def test_export_returns_bytesio(self):
        """Verify export() returns BytesIO with valid DOCX content (ZIP magic bytes)."""
        from unittest.mock import MagicMock

        from exporters.docx import DocxExporter

        # Create mock records
        record = MagicMock()
        record.content = '<p>Test content</p>'
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = DocxExporter(uploads_path='/tmp')
        result = exporter.export([record], title='Test Report')

        assert isinstance(result, BytesIO)
        # DOCX files are ZIP archives with PK magic bytes
        result.seek(0)
        header = result.read(2)
        assert header == b'PK'

    def test_html_to_docx_conversion(self):
        """Verify HTML elements are converted to DOCX equivalents."""
        from unittest.mock import MagicMock

        from docx import Document

        from exporters.docx import DocxExporter

        # Create mock record with various HTML elements
        record = MagicMock()
        record.content = '''
        <h1>Heading 1</h1>
        <h2>Heading 2</h2>
        <h3>Heading 3</h3>
        <p>Paragraph with <strong>bold</strong> and <em>italic</em> text.</p>
        <ul>
            <li>Unordered item 1</li>
            <li>Unordered item 2</li>
        </ul>
        <ol>
            <li>Ordered item 1</li>
            <li>Ordered item 2</li>
        </ol>
        <table>
            <tr><th>Header</th></tr>
            <tr><td>Cell</td></tr>
        </table>
        <a href="https://example.com">Link</a>
        <pre><code>Code block</code></pre>
        '''
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = DocxExporter(uploads_path='/tmp')
        result = exporter.export([record], title='Test Report')

        # Parse the DOCX to verify structure
        result.seek(0)
        from docx import Document
        doc = Document(result)

        # Verify document is not empty
        assert len(doc.paragraphs) > 0

        # Verify headings exist (document should have styled paragraphs)
        heading_found = any(p.style.name.startswith('Heading') for p in doc.paragraphs)
        assert heading_found, "Document should contain heading styles"

        # Verify lists exist (bullet points or numbered)
        # Lists in python-docx appear as paragraphs with list styling
        list_found = any(
            p.style.name in ['List Bullet', 'List Number', 'List Paragraph']
            for p in doc.paragraphs
        )
        assert list_found, "Document should contain list styles"

        # Verify tables exist
        assert len(doc.tables) > 0, "Document should contain tables"

    def test_image_embedding(self):
        """Verify image embedding via _extract_images and _add_image_to_document methods."""
        import os
        import tempfile
        from unittest.mock import MagicMock

        from docx import Document
        from PIL import Image

        from exporters.docx import DocxExporter

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test image file (valid 1x1 pixel PNG using PIL)
            test_image = os.path.join(tmpdir, 'test.png')
            img = Image.new('RGB', (1, 1), color='red')
            img.save(test_image, 'PNG')

            # Create mock record with image
            record = MagicMock()
            record.content = '<p>Text before image.</p><img src="/files/test.png"><p>Text after image.</p>'
            record.date.strftime = lambda fmt: '2026-03-26'

            exporter = DocxExporter(uploads_path=tmpdir)
            result = exporter.export([record], title='Test Report')

            # Parse the DOCX and verify image is embedded
            result.seek(0)
            doc = Document(result)

            # Check that document has inline shapes (images)
            # Images in python-docx appear as InlineShape objects
            image_count = 0
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        image_count += 1
                    # Also check for drawing elements
                    if run._element.xpath('.//w:drawing'):
                        image_count += 1

            # Verify at least one image was found
            assert image_count > 0, "Document should contain embedded images"

    def test_extract_images_helper(self):
        """Verify _extract_images extracts img tags from HTML."""
        from exporters.docx import DocxExporter

        exporter = DocxExporter(uploads_path='/tmp')

        html = '''
        <p>Text before</p>
        <img src="/files/image1.png">
        <p>Text between</p>
        <img src="/files/image2.jpg" alt="Second image">
        <p>Text after</p>
        '''

        result = exporter._extract_images(html)

        # Should return tuple of (processed_html, images_list)
        assert isinstance(result, tuple)
        processed_html, images = result

        # Check images list
        assert isinstance(images, list)
        assert len(images) == 2
        assert images[0][1] == '/files/image1.png'
        assert images[1][1] == '/files/image2.jpg'

        # Check that placeholders were inserted in HTML
        assert '[[IMAGE_PLACEHOLDER_0]]' in processed_html
        assert '[[IMAGE_PLACEHOLDER_1]]' in processed_html

    def test_add_image_to_document_helper(self):
        """Verify _add_image_to_document embeds image bytes into document."""
        import os
        import tempfile

        from docx import Document
        from PIL import Image

        from exporters.docx import DocxExporter

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test image using PIL
            test_image = os.path.join(tmpdir, 'embed_test.png')
            img = Image.new('RGB', (1, 1), color='blue')
            img.save(test_image, 'PNG')

            exporter = DocxExporter(uploads_path=tmpdir)
            doc = Document()

            # Add a paragraph first
            paragraph = doc.add_paragraph()

            # Add image to document
            image_bytes = exporter.image_resolver.get_image_bytes('/files/embed_test.png')
            assert image_bytes is not None, "Image bytes should be retrieved"

            exporter._add_image_to_document(doc, paragraph, '/files/embed_test.png', image_bytes)

            # Verify image was added (document should have relationships)
            # Check for drawing elements in the paragraph
            has_drawing = len(paragraph.runs) > 0 and any(
                run._element.xpath('.//w:drawing') for run in paragraph.runs
            )
            assert has_drawing, "Paragraph should contain drawing element after image added"


class TestExcelExporter:
    """Test ExcelExporter class behavior.

    Tests for Excel export with rich text functionality to be implemented in Phase 11 Plan 01.
    These tests are designed to FAIL initially and will pass once ExcelExporter is implemented.
    """

    def test_file_extension(self):
        """Verify ExcelExporter.file_extension returns 'xlsx'."""
        from exporters.excel import ExcelExporter
        exporter = ExcelExporter()
        assert exporter.file_extension == 'xlsx'

    def test_mime_type(self):
        """Verify ExcelExporter.mime_type returns correct XLSX MIME type."""
        from exporters.excel import ExcelExporter
        exporter = ExcelExporter()
        expected_mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        assert exporter.mime_type == expected_mime

    def test_export_returns_bytesio(self):
        """Verify export() returns BytesIO with valid XLSX content (ZIP magic bytes)."""
        from unittest.mock import MagicMock

        from exporters.excel import ExcelExporter

        # Create mock records
        record = MagicMock()
        record.content = '<p>Test content</p>'
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = ExcelExporter()
        result = exporter.export([record], title='Test Report')

        assert isinstance(result, BytesIO)
        # XLSX files are ZIP archives with PK magic bytes
        result.seek(0)
        header = result.read(2)
        assert header == b'PK'

    def test_html_to_rich_text_bold(self):
        """Verify _html_to_rich_text converts <strong> to CellRichText with bold InlineFont."""
        from openpyxl.cell.rich_text import CellRichText, TextBlock

        from exporters.excel import ExcelExporter

        exporter = ExcelExporter()
        result = exporter._html_to_rich_text('<strong>Bold text</strong>')

        assert isinstance(result, CellRichText)
        # Should have at least one TextBlock with bold font
        assert len(result) >= 1
        # Check that the first text block has bold font
        first_block = result[0]
        assert isinstance(first_block, TextBlock)
        assert first_block.font.b is True

    def test_html_to_rich_text_italic(self):
        """Verify _html_to_rich_text converts <em> to CellRichText with italic InlineFont."""
        from openpyxl.cell.rich_text import CellRichText, TextBlock

        from exporters.excel import ExcelExporter

        exporter = ExcelExporter()
        result = exporter._html_to_rich_text('<em>Italic text</em>')

        assert isinstance(result, CellRichText)
        assert len(result) >= 1
        first_block = result[0]
        assert isinstance(first_block, TextBlock)
        assert first_block.font.i is True

    def test_html_to_rich_text_underline(self):
        """Verify _html_to_rich_text converts <u> to CellRichText with underline='single'."""
        from openpyxl.cell.rich_text import CellRichText, TextBlock

        from exporters.excel import ExcelExporter

        exporter = ExcelExporter()
        result = exporter._html_to_rich_text('<u>Underline text</u>')

        assert isinstance(result, CellRichText)
        assert len(result) >= 1
        first_block = result[0]
        assert isinstance(first_block, TextBlock)
        # Note: underline must be 'single' string, not boolean True
        assert first_block.font.u == 'single'

    def test_html_to_rich_text_nested(self):
        """Verify _html_to_rich_text handles nested formatting like <strong>Bold <em>and italic</em></strong>."""
        from openpyxl.cell.rich_text import CellRichText, TextBlock

        from exporters.excel import ExcelExporter

        exporter = ExcelExporter()
        result = exporter._html_to_rich_text('<strong>Bold <em>and italic</em></strong>')

        assert isinstance(result, CellRichText)
        # Should have multiple text blocks for nested formatting
        assert len(result) >= 2
        # First block should be bold only
        first_block = result[0]
        assert isinstance(first_block, TextBlock)
        assert first_block.font.b is True
        assert first_block.font.i is False
        # Second block should be bold+italic
        second_block = result[1]
        assert isinstance(second_block, TextBlock)
        assert second_block.font.b is True
        assert second_block.font.i is True

    def test_rich_text_in_cell(self):
        """Verify exported cell value is CellRichText instance when HTML formatting present."""
        from unittest.mock import MagicMock

        from openpyxl.cell.rich_text import CellRichText

        from exporters.excel import ExcelExporter

        # Create mock record with rich text
        record = MagicMock()
        record.content = '<p><strong>Bold</strong> and <em>italic</em></p>'
        record.date.strftime = lambda fmt: '2026-03-26'

        exporter = ExcelExporter()
        result = exporter.export([record], title='Test Report')

        # Re-open to check cell content
        from openpyxl import load_workbook
        result.seek(0)
        wb = load_workbook(result, rich_text=True)  # rich_text=True preserves CellRichText
        ws = wb.active

        # Find a cell with rich text content (skip header row)
        found_rich_text = False
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                if isinstance(cell.value, CellRichText):
                    found_rich_text = True
                    break
            if found_rich_text:
                break

        assert found_rich_text, "At least one cell should contain CellRichText"
