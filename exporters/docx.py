"""DOCX exporter using python-docx with htmldocx for HTML conversion.

This module provides the DocxExporter class for generating DOCX documents
from weekly report records with embedded images and Chinese font support.
"""

from docx import Document
from docx.shared import Inches, Pt
from htmldocx import HtmlToDocx
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
import os
from datetime import datetime
from urllib.parse import unquote

from .base import ExporterBase
from .image_resolver import ImageResolver


class DocxExporter(ExporterBase):
    """DOCX exporter using python-docx with htmldocx for HTML conversion.

    Generates DOCX documents from weekly report records with:
    - Embedded images from /files/ URLs (custom handling required)
    - Chinese font support (Microsoft YaHei)
    - HTML formatting preserved (headings, lists, tables, links, code)

    Attributes:
        _uploads_path: Path to uploads directory (injected or from Flask config)
        _image_resolver: Lazy-loaded ImageResolver instance
    """

    def __init__(self, uploads_path: Optional[str] = None):
        """Initialize with optional uploads path for dependency injection.

        Args:
            uploads_path: Path to uploads directory. If None, will be
                         initialized from Flask current_app.config when needed.
        """
        self._uploads_path = uploads_path
        self._image_resolver: Optional[ImageResolver] = None

    @property
    def uploads_path(self) -> str:
        """Get uploads path, initializing from Flask config if needed."""
        if self._uploads_path is None:
            from flask import current_app
            self._uploads_path = current_app.config['UPLOADED_PATH']
        return self._uploads_path

    @property
    def image_resolver(self) -> ImageResolver:
        """Get image resolver instance (lazy initialization)."""
        if self._image_resolver is None:
            self._image_resolver = ImageResolver(self.uploads_path)
        return self._image_resolver

    @property
    def file_extension(self) -> str:
        """Return file extension without dot."""
        return 'docx'

    @property
    def mime_type(self) -> str:
        """Return MIME type for send_file()."""
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    def _generate(self, records: List[Any], options: Dict) -> BytesIO:
        """Generate DOCX from records using python-docx.

        Args:
            records: List of Record objects with content and date
            options: Export options (title, include_date, etc.)

        Returns:
            BytesIO buffer containing DOCX
        """
        title = options.get('title', 'Weekly Report')

        # Create document
        doc = Document()

        # Add title
        doc.add_heading(title, level=0)

        # Process each record
        for record in records:
            if record.content:
                # Get user names for this record
                user_names = ', '.join(u.username for u in record.user) if record.user else '未知用户'
                date_str = record.date.strftime('%Y-%m-%d')

                # Add user and date paragraph
                doc.add_paragraph(f'{user_names} — {date_str}', style='Intense Quote')

                # Convert HTML to DOCX
                self._convert_html_to_document(doc, record.content)

                # Add separator
                doc.add_paragraph('---')

        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    def _convert_html_to_document(self, doc: Document, html: str) -> None:
        """Convert HTML content to DOCX document elements.

        Uses htmldocx for HTML conversion and applies Chinese font support.

        Args:
            doc: python-docx Document object
            html: HTML content string
        """
        # First, extract and embed images (htmldocx doesn't support images)
        html_processed, images = self._extract_images(html)

        # Pre-process HTML to fix issues with htmldocx
        html_processed = self._sanitize_html(html_processed)

        # Convert HTML to DOCX using htmldocx
        parser = HtmlToDocx()
        parser.add_html_to_document(html_processed, doc)

        # Apply Chinese font to all runs
        self._apply_chinese_font(doc)

        # Embed extracted images
        self._embed_images(doc, images)

    def _sanitize_html(self, html: str) -> str:
        """Sanitize HTML to avoid htmldocx parsing issues.

        htmldocx has bugs with:
        - <a> tags without href attribute (KeyError: 'href')

        Args:
            html: HTML content string

        Returns:
            Sanitized HTML string
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, 'html.parser')

        # Fix <a> tags without href - add a dummy href or convert to span
        for a in soup.find_all('a'):
            if not a.get('href'):
                # Convert to span to avoid htmldocx KeyError
                a.name = 'span'

        return str(soup)

    def _extract_images(self, html: str) -> Tuple[str, list]:
        """Extract img tags from HTML and replace with placeholders.

        htmldocx does NOT support images, so we extract them first,
        replace with placeholder text, then embed them via python-docx.

        Args:
            html: HTML content string

        Returns:
            Tuple of (processed_html, list of (placeholder, url) tuples)
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, 'html.parser')
        images = []

        for i, img in enumerate(soup.find_all('img')):
            src = img.get('src', '')
            if src:
                placeholder = f'[[IMAGE_PLACEHOLDER_{i}]]'
                images.append((placeholder, src))
                img.replace_with(soup.new_string(placeholder))

        return str(soup), images

    def _embed_images(self, doc: Document, images: list) -> None:
        """Embed extracted images into the document.

        Finds placeholder text in document and replaces with actual images.

        Args:
            doc: python-docx Document object
            images: List of (placeholder, url) tuples
        """
        for placeholder, url in images:
            # Get image bytes via ImageResolver
            image_bytes = self.image_resolver.get_image_bytes(url)

            if image_bytes:
                # Find paragraph containing the placeholder
                found = False
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        # Clear the paragraph and add image
                        self._replace_paragraph_with_image(paragraph, image_bytes)
                        found = True
                        break

                # If placeholder not found, append image at end
                if not found:
                    self._add_image_to_document(doc, None, url, image_bytes)

    def _replace_paragraph_with_image(self, paragraph, image_bytes: bytes) -> None:
        """Replace paragraph content with an image.

        Args:
            paragraph: python-docx Paragraph object
            image_bytes: Image data as bytes
        """
        from io import BytesIO as IOBytesIO

        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''

        # Add image to the paragraph
        image_stream = IOBytesIO(image_bytes)
        try:
            paragraph.add_run().add_picture(image_stream, width=Inches(6))
        except Exception:
            paragraph.add_run('[Image]')

    def _add_image_to_document(self, doc: Document, paragraph, url: str, image_bytes: bytes) -> None:
        """Add image bytes to document at appropriate location.

        Args:
            doc: python-docx Document object
            paragraph: Paragraph to add image to (or None to append new)
            url: Original image URL (for logging)
            image_bytes: Image data as bytes
        """
        from io import BytesIO as IOBytesIO

        # Create BytesIO from image bytes for python-docx
        image_stream = IOBytesIO(image_bytes)

        # Add image to document
        # Use a reasonable default width (6 inches = ~15cm)
        try:
            if paragraph:
                paragraph.add_run().add_picture(image_stream, width=Inches(6))
            else:
                # Add new paragraph with image
                doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(6))
        except Exception:
            # If image embedding fails, add placeholder text
            if paragraph:
                paragraph.add_run('[Image]')
            else:
                doc.add_paragraph('[Image]')

    def _apply_chinese_font(self, doc: Document) -> None:
        """Apply Microsoft YaHei font to all runs in document for Chinese support.

        Args:
            doc: python-docx Document object
        """
        from docx.oxml.ns import qn

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Microsoft YaHei'
                # Also set East Asian font via rPr
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')